import docx
from pdfminer.high_level import extract_text as pdf_extract_text
import io

def extract_text_from_file(file_storage):
    filename = file_storage.filename.lower()
    warning = None

    if filename.endswith('.txt'):
        text = file_storage.read().decode('utf-8', errors='ignore')

    elif filename.endswith('.docx'):
        doc = docx.Document(file_storage)
        text = "\n".join([p.text for p in doc.paragraphs])

    elif filename.endswith('.pdf'):
        try:
            text = pdf_extract_text(file_storage)
            if not text.strip():
                warning = "PDF appears to be scanned or unreadable. Please use Word format."
        except Exception:
            text = ""
            warning = "Error reading PDF. Please use Word format."
    else:
        text = ""
        warning = "Unsupported file type."

    return text, warning


def create_docx_from_text(translated_text):
    doc = docx.Document()
    for paragraph in translated_text.split('\n'):
        doc.add_paragraph(paragraph)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
